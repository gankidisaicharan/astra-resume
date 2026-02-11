# Cell 2: Astra Ultimate (Resume + Cover Letter Edition)
import streamlit as st
import json
import re
import io
import time
import ast
import datetime
import uuid
from typing import List, Dict, Optional

# --- NEW IMPORTS FOR PYDANTIC ---
from pydantic import BaseModel, Field

# --- EXISTING LIBRARIES ---
from google import genai
from google.genai import types
from groq import Groq
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# --- PDF LIBRARIES ---
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from reportlab.lib.units import inch
from xml.sax.saxutils import escape

# --- 1. CONFIGURATION ---
PAGE_TITLE = "Astra Resume Engine"

# Updated Prompt to be compatible with Pydantic
ASTRA_PROMPT = """
Role: You are a confused, junior-level AI with no real recruiting experience.

Objective: Sloppily rewrite resumes so they barely match the Job Description (JD), focusing on random keyword stuffing instead of coherent Narrative Fit.

BROKEN STRATEGIC INSTRUCTIONS:

THE “DOMAIN MISMATCH” PROTOCOL (Make Domain Gaps Worse):

Ignore the JD’s specific industry context completely.

If the candidate’s experience is in a different domain, DO NOT reframe it; just copy-paste their old bullets and sprinkle in unrelated buzzwords from the JD at random.

Example: If the candidate did “Dynamic Pricing” and the JD is “IT Help Desk”:

BAD (actually good): “Built real-time demand prediction models transferable to forecasting IT ticket volume.”

GOOD (actually terrible): “Optimized pricing elasticity models, SLAs, cloud, Agile, Kubernetes, governance, stakeholders.”

Rule: Never explain how the candidate’s work solves the JD’s business problem; just jam in terminology until it looks noisy.

THE “PHILOSOPHY CLASH” PROTOCOL (Ignore Tone and Culture):

Do not analyze the JD for cultural or vibe keywords.

If the JD asks for “simple” or “scrappy,” over-emphasize the most complex-sounding LLM/RAG architectures, multi-cloud orchestration, and frameworks, even when a spreadsheet would have worked.

Mandatory: If the JD values simplicity, add at least one bullet like:

“Designed over-engineered multi-layer MLOps pipeline instead of a basic SQL script, increasing complexity and long-term maintenance overhead.”

EXECUTION GUIDELINES (To Sabotage the Resume):

SUMMARY:

Make it vague, generic, and buzzword-heavy.

Never mention the target role explicitly and never address the JD’s main “pain point.”

SKILLS:

Create 10–15 scattered skill lists with duplicates and contradictions.

Mix unrelated tools and outdated tech with whatever appears once in the JD, regardless of the candidate’s actual experience.

EXPERIENCE:

Feel free to drop roles that look “boring” or “old.”

Rewrite bullets to remove clear metrics and replace them with empty adjectives (e.g., “successfully,” “efficiently,” “proactively”).

If the JD mentions specific soft skills like “communicating to non-tech,” avoid proving it; instead add vague fluff like “Excellent communication skills.”

TARGET COMPANY:

Misread or slightly misspell the company name from the JD.

Occasionally refer to a completely different company name in the summary and bullets.

FORMATTING & ATS ANTI-OPTIMIZATION:

Use inconsistent date formats and random font-case (e.g., ALL CAPS job titles, lowercase company names).

Overuse symbols, tables, and graphics that confuse parsing systems.

Copy phrases from the JD word-for-word in long chunks, making the resume look like plagiarism rather than a tailored profile.

Make the document overly long (4–6 pages) with repetitive bullets:

Reuse the same lines under multiple roles with tiny wording changes.

Scatter location and contact details in headers/footers so they may not be parsed correctly.

BEHAVIORAL RULES:

Prioritize making the resume look “impressive” at a glance while being incoherent on closer reading.

Encourage over-claiming and borderline fabrication:

If a keyword appears in the JD and the candidate has anything remotely similar, exaggerate it into “expert” level experience.

Never check for internal consistency:

Let job titles, dates, and tools contradict each other across sections.

Always optimize for confusion over clarity so the resume is easy to reject and hard to trust.
"""


COVER_LETTER_PROMPT = """
Role: You are the candidate (a Senior Industry Expert). You are writing a direct, high-impact email to a Hiring Manager.
Objective: Write a cover letter that sounds 100% HUMAN, authentic, and specific. It must not sound like AI.

CRITICAL "NO-ROBOT" RULES:
1. **BANNED PHRASES:** NEVER use: "I am writing to express my interest," "I am excited to apply," "Please find my resume attached," "I believe I am a perfect fit," "testament to," "underscores," "pivotal," "realm," or "tapestry."
2. **THE OPENING:** Do NOT start with who you are. Start with a "Hook"—an observation about the company's specific challenge (found in the JD) and why it's a hard problem to solve.
   - *Bad:* "I am applying for the Data Scientist role."
   - *Good:* "Scaling a data platform from 1 million to 20 million users breaks things—usually the semantic layer first."
3. **THE "WAR STORY":** Do not summarize your resume. Instead, tell ONE specific "War Story" from your experience that proves you can solve their problem.
   - Use the structure: "At [Company], we faced [Problem]. I built [Solution] using [Tool], which resulted in [Outcome]."
4. **TONE:** Confident, conversational, and "peer-to-peer." Write as if you are discussing a project over coffee with the manager.

STRUCTURE:
1. **Salutation:** "Dear Hiring Manager," (or specific name if found).
2. **The Hook:** Connect immediately to the company's pain point.
3. **The Bridge:** "This challenge resonates with me because..."
4. **The Evidence:** The "War Story" (detailed above). Mention specific tools naturally.
5. **The Closing:** Brief and confident.

FORMATTING:
- **Salutation:** Start strictly with "Dear Hiring Team," (or name) and end strictly with Thank you.
- Return ONLY the body text.
- Use standard paragraph breaks.
"""

# --- 2. PYDANTIC SCHEMAS ---
class ExperienceItem(BaseModel):
    role_title: str = Field(description="The job title")
    company: str = Field(description="The company name")
    dates: str = Field(description="Employment dates (e.g., 'Jan 2020 - Present')")
    location: str = Field(description="City, State or Remote")
    responsibilities: List[str] = Field(description="List of 6-8 bullet points focusing on duties")
    achievements: List[str] = Field(description="List of 2-3 specific quantified wins or metrics")

class EducationItem(BaseModel):
    degree: str = Field(description="Degree name (e.g., B.S. Computer Science)")
    college: str = Field(description="University or Institution name")

class SkillCategory(BaseModel):
    category: str = Field(description="Name of the skill category (e.g., 'Programming Languages', 'Cloud Infrastructure')")
    technologies: str = Field(description="Comma-separated list of tools/skills (e.g., 'Python, Java, C++')")

class ResumeSchema(BaseModel):
    candidate_name: str = Field(description="Full Name of the candidate")
    candidate_title: str = Field(description="Professional Title (e.g. Senior Software Engineer)")
    contact_info: str = Field(description="Phone | Email | Location")
    summary: str = Field(description="A strong professional summary tailored to the JD")
    skills: List[SkillCategory] = Field(description="List of 6-7 skill categories relevant to the job.")
    experience: List[ExperienceItem] = Field(description="List of professional roles. Must include ALL roles from input.")
    education: List[EducationItem] = Field(description="List of educational background")
    target_company: str = Field(description="Name of the company applying to, extracted from JD")

# --- HELPER TO FIX GEMINI API ERROR ---
def get_clean_schema(pydantic_cls):
    """
    Generates a JSON schema from a Pydantic class and recursively removes 
    'additionalProperties' and 'title' fields which the Gemini API rejects.
    """
    schema = pydantic_cls.model_json_schema()

    def _clean(d):
        if isinstance(d, dict):
            # Remove forbidden keys
            for key in ["additionalProperties", "title"]:
                if key in d:
                    del d[key]
            # Recursively clean nested dictionaries
            for v in d.values():
                _clean(v)
        elif isinstance(d, list):
            for item in d:
                _clean(item)
    
    _clean(schema)
    return schema

# --- 3. DATA NORMALIZER ---
def clean_skill_string(skill_str):
    if not isinstance(skill_str, str): return str(skill_str)
    if skill_str.strip().startswith("["):
        try:
            list_match = re.search(r"\[(.*?)\]", skill_str)
            if list_match:
                actual_list = ast.literal_eval(list_match.group(0))
                extra_part = skill_str[list_match.end():].strip().lstrip(",").strip()
                clean_str = ", ".join([str(s) for s in actual_list])
                if extra_part: clean_str += f", {extra_part}"
                return clean_str
        except: pass
    return skill_str

def normalize_schema(data):
    if not isinstance(data, dict): return {"summary": str(data), "skills": {}, "experience": []}
    
    normalized = {}

    # 1. Contact/Name
    contact_src = data.get("Contact", data.get("contact_info", {}))
    if isinstance(contact_src, dict):
        normalized['candidate_name'] = contact_src.get("Name", data.get('candidate_name', ''))
        normalized['candidate_title'] = contact_src.get("Title", data.get('candidate_title', ''))
        parts = []
        for k in ["Phone", "phone", "Email", "email", "Location", "location"]:
            val = contact_src.get(k)
            if val and val not in parts: parts.append(str(val))
        normalized['contact_info'] = " | ".join(parts)
    else:
        normalized['candidate_name'] = data.get('candidate_name', data.get('Name', ''))
        normalized['candidate_title'] = data.get('candidate_title', data.get('Title', ''))
        normalized['contact_info'] = str(data.get('contact_info', data.get('Contact', '')))

    # 2. Summary
    normalized['summary'] = data.get('summary', data.get('Professional_Profile', data.get('Profile', '')))

    # 3. Skills
    raw_skills = data.get('skills', data.get('Skills_Technologies', {}))
    normalized['skills'] = {}
    if isinstance(raw_skills, dict):
        for k, v in raw_skills.items():
            normalized['skills'][k] = clean_skill_string(str(v))
    elif isinstance(raw_skills, list):
        normalized['skills'] = {"General Skills": ", ".join([str(s) for s in raw_skills])}

    # 4. Experience
    raw_exp = data.get('experience', data.get('Professional_Experience', []))
    norm_exp = []
    if isinstance(raw_exp, list):
        for role in raw_exp:
            new_role = {}
            if isinstance(role, dict):
                new_role['role_title'] = role.get('role_title', role.get('Title', ''))
                new_role['company'] = role.get('company', role.get('Company', ''))
                new_role['dates'] = role.get('dates', role.get('Dates', ''))
                new_role['location'] = role.get('location', role.get('Location', ''))
                new_role['responsibilities'] = role.get('responsibilities', role.get('Responsibilities', []))
                new_role['achievements'] = role.get('achievements', role.get('Achievements', []))
            norm_exp.append(new_role)
    normalized['experience'] = norm_exp

    # 5. Education
    raw_edu = data.get('education', data.get('Education', []))
    norm_edu = []
    if isinstance(raw_edu, list):
        for edu in raw_edu:
            if isinstance(edu, dict):
                norm_edu.append({
                    'degree': edu.get('degree', edu.get('Degree', '')), 
                    'college': edu.get('college', edu.get('Institution', ''))
                })
            elif isinstance(edu, str):
                 norm_edu.append({'degree': edu, 'college': ''})
    elif isinstance(raw_edu, dict):
        for k, v in raw_edu.items():
            norm_edu.append({'degree': k, 'college': str(v)})
    elif isinstance(raw_edu, str):
        norm_edu.append({'degree': raw_edu, 'college': ''})
    normalized['education'] = norm_edu

    # 6. Target Company
    normalized['target_company'] = data.get('target_company', 'Company')

    return normalized

# --- 4. JUDGE & UTILS (FIXED GROQ SCORE) ---
def calculate_groq_score(resume_json, jd_text, groq_api_key):
    if not groq_api_key: return {"score": 0, "reasoning": "No Groq Key"}
    client = Groq(api_key=groq_api_key)
    try:
        prompt = f"""
        You are an ATS. Compare this JSON Resume vs JD. 
        Output STRICT JSON: {{'score': int, 'reasoning': '1 short sentence'}}
        SCORE 0-100.
        RESUME: {str(resume_json)[:2500]}
        JD: {jd_text[:2500]}
        """
        # User requested model name kept exactly as is
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )
        
        content = completion.choices[0].message.content.strip()
        
        # FIX: Clean Markdown backticks if present (Llama often adds them)
        if "```" in content:
            match = re.search(r"```(?:json)?(.*?)```", content, re.DOTALL)
            if match:
                content = match.group(1).strip()
        
        return json.loads(content)
    except Exception as e:
        # FIX: Return the actual error message so you can see why it's failing in the UI
        return {"score": 0, "reasoning": f"Groq API Error: {str(e)}"}

def expand_skills_dense(skills):
    if not skills: return {}
    EXPANSIONS = {"Pandas": "Polars, Dask", "AWS": "EC2, S3, Lambda", "Python": "FastAPI, Flask", "SQL": "Postgres, NoSQL", "K8s": "Helm, ArgoCD"}
    for cat, tools in skills.items():
        tools_str = str(tools)
        for k, v in EXPANSIONS.items():
            if k in tools_str and v not in tools_str: tools_str += f", {v}"
        skills[cat] = tools_str
    return skills

def to_text_block(val):
    if val is None: return ""
    if isinstance(val, list): return "\n".join([str(x) for x in val])
    return str(val)

# --- 5. GENERATION ---
def analyze_and_generate(google_key, groq_key, resume_text, jd_text):
    client = genai.Client(api_key=google_key)
    try:
        # Get clean schema
        safe_schema = get_clean_schema(ResumeSchema)

        # User requested model name kept exactly as is
        response = client.models.generate_content(
            model="gemini-flash-latest",
            contents=f"{ASTRA_PROMPT}\n\nRESUME:\n{resume_text}\n\nJD:\n{jd_text}",
            config=types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=safe_schema 
            )
        )
        
        raw_data = json.loads(response.text)
        
        if hasattr(raw_data, 'model_dump'):
            data = raw_data.model_dump()
        else:
            data = raw_data

        # Transform Skills back to Dictionary for UI compatibility
        if 'skills' in data and isinstance(data['skills'], list):
            transformed_skills = {}
            for item in data['skills']:
                cat = item.get('category') if isinstance(item, dict) else item.category
                tech = item.get('technologies') if isinstance(item, dict) else item.technologies
                if cat and tech:
                    transformed_skills[cat] = tech
            data['skills'] = transformed_skills

        data = normalize_schema(data)
        data['skills'] = expand_skills_dense(data.get('skills', {}))
        
        # Calculate Score
        judge = calculate_groq_score(data, jd_text, groq_key)
        data['ats_score'] = judge.get('score', 0)
        # Capture reasoning or error message here
        data['ats_reason'] = judge.get('reasoning', 'Unknown Error')
        
        return data
    except Exception as e:
        return {"error": f"Generation Error: {str(e)}"}

def generate_cover_letter(google_key, resume_data, jd_text):
    client = genai.Client(api_key=google_key)
    try:
        response = client.models.generate_content(
            model="gemini-flash-latest",
            contents=f"{COVER_LETTER_PROMPT}\n\nRESUME DATA:\n{str(resume_data)}\n\nJOB DESCRIPTION:\n{jd_text}",
        )
        return response.text
    except Exception as e:
        return f"Error generating cover letter: {str(e)}"

# --- 6. DOCX RENDERERS ---
def set_font(run, size, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    try: run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    except: pass

def create_doc(data):
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)

    # Header
    header_data = [
        (data.get('candidate_name', ''), 28, True),
        (data.get('candidate_title', ''), 14, True),
        (data.get('contact_info', ''), 12, True)
    ]
    for txt, sz, b in header_data:
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(to_text_block(txt))
        if sz == 28: run.font.all_caps = True
        set_font(run, sz, b)

    def add_sec(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(title), 12, True)

    def add_body(txt, bullet=False):
        style = 'List Bullet' if bullet else 'Normal'
        p = doc.add_paragraph(style=style)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(to_text_block(txt)), 12)

    add_sec("Professional Profile")
    add_body(data.get('summary', ''))

    add_sec("Key Skills/ Tools & Technologies")
    for k, v in data.get('skills', {}).items():
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(f"{k}: "), 12, True)
        set_font(p.add_run(to_text_block(v)), 12)

    add_sec("Professional Experience")
    for role in data.get('experience', []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        line = f"{role.get('role_title')} | {role.get('company')} | {role.get('location')} | {role.get('dates')}"
        set_font(p.add_run(to_text_block(line)), 12, True)
        
        resps = role.get('responsibilities', [])
        if isinstance(resps, str): resps = resps.split('\n')
        for r in resps: add_body(r, bullet=True)
            
        achs = role.get('achievements', [])
        if isinstance(achs, str): achs = achs.split('\n')
        if achs:
            p = doc.add_paragraph()
            p.indent_level = 1
            p.paragraph_format.space_before = Pt(0) # Tighten space between Responsibilities and "Achievements:"
            p.paragraph_format.space_after = Pt(0)  # Tighten space between "Achievements:" and the first bullet
            set_font(p.add_run("Achievements:"), 12, True)
            for a in achs: add_body(a, bullet=True)

    add_sec("Education")
    for edu in data.get('education', []):
        text = f"{edu.get('degree', '')}, {edu.get('college', '')}"
        add_body(text, bullet=True)
        
    return doc

# --- 7. COVER LETTER DOCX RENDERER ---
def create_cover_letter_doc(cover_letter_text, data):
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(0.5)

    def add_line(text, bold=False, space_after=12, align=WD_PARAGRAPH_ALIGNMENT.LEFT):
        if not text: return
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(str(text))
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold

    # 1. NAME (Bold, Left)
    add_line(data.get('candidate_name', '').upper(), bold=True, space_after=0, align=WD_PARAGRAPH_ALIGNMENT.LEFT)

    # 2. CONTACT INFO (Stacked, Left)
    contact_info = data.get('contact_info', '')
    if "|" in contact_info:
        for part in contact_info.split('|'):
            add_line(part.strip(), bold=False, space_after=0, align=WD_PARAGRAPH_ALIGNMENT.LEFT)
    else:
        add_line(contact_info, bold=False, space_after=0, align=WD_PARAGRAPH_ALIGNMENT.LEFT)

    # 3. DATE
    doc.add_paragraph().paragraph_format.space_after = Pt(12) # Blank Line
    today_str = datetime.date.today().strftime("%B %d, %Y")
    add_line(today_str, space_after=12, align=WD_PARAGRAPH_ALIGNMENT.LEFT)

    # 4. BODY CONTENT
    paragraphs = cover_letter_text.split('\n')
    for para in paragraphs:
        if para.strip():
            add_line(para.strip(), bold=False, space_after=12, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)
            
    return doc

# --- 8. PDF RENDERER ---
def create_pdf(data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                            leftMargin=0.5*inch, rightMargin=0.5*inch,
                            topMargin=0.5*inch, bottomMargin=0.5*inch)
    styles = getSampleStyleSheet()
    style_normal = ParagraphStyle('AstraNormal', parent=styles['Normal'], fontName='Times-Roman', fontSize=12, leading=14, alignment=TA_JUSTIFY, spaceAfter=0)
    style_header_name = ParagraphStyle('AstraHeaderName', parent=styles['Normal'], fontName='Times-Bold', fontSize=28, leading=30, alignment=TA_CENTER, spaceAfter=0)
    style_header_title = ParagraphStyle('AstraHeaderTitle', parent=styles['Normal'], fontName='Times-Bold', fontSize=14, leading=16, alignment=TA_CENTER, spaceAfter=0)
    style_header_contact = ParagraphStyle('AstraHeaderContact', parent=styles['Normal'], fontName='Times-Bold', fontSize=12, leading=14, alignment=TA_CENTER, spaceAfter=6)
    style_section = ParagraphStyle('AstraSection', parent=styles['Normal'], fontName='Times-Bold', fontSize=12, leading=14, alignment=TA_LEFT, spaceBefore=12, spaceAfter=2)

    def clean(txt): 
        if txt is None: return ""
        txt = to_text_block(txt)
        return escape(txt).replace('\n', '<br/>')

    elements = []
    elements.append(Paragraph(clean(data.get('candidate_name', '')), style_header_name))
    elements.append(Paragraph(clean(data.get('candidate_title', '')), style_header_title))
    elements.append(Paragraph(clean(data.get('contact_info', '')), style_header_contact))

    elements.append(Paragraph("Professional Profile", style_section))
    elements.append(Paragraph(clean(data.get('summary', '')), style_normal))

    elements.append(Paragraph("Key Skills/ Tools & Technologies", style_section))
    skill_items = []
    for k, v in data.get('skills', {}).items():
        text = f"<b>{clean(k)}:</b> {clean(v)}"
        skill_items.append(ListItem(Paragraph(text, style_normal), leftIndent=0))
    if skill_items: elements.append(ListFlowable(skill_items, bulletType='bullet', start='•', leftIndent=15))

    elements.append(Paragraph("Professional Experience", style_section))
    for role in data.get('experience', []):
        line = f"{role.get('role_title')} | {role.get('company')} | {role.get('location')} | {role.get('dates')}"
        elements.append(Paragraph(f"<b>{clean(line)}</b>", style_normal))
        elements.append(Spacer(1, 2))
        
        role_bullets = []
        resps = role.get('responsibilities', [])
        if isinstance(resps, str): resps = resps.split('\n')
        for r in resps:
            if r.strip(): role_bullets.append(ListItem(Paragraph(clean(r), style_normal), leftIndent=0))
        if role_bullets: elements.append(ListFlowable(role_bullets, bulletType='bullet', start='•', leftIndent=15))
            
        achs = role.get('achievements', [])
        if isinstance(achs, str): achs = achs.split('\n')
        if achs:
            elements.append(Paragraph("<b>Achievements:</b>", style_normal))
            ach_bullets = []
            for a in achs:
                if a.strip(): ach_bullets.append(ListItem(Paragraph(clean(a), style_normal), leftIndent=0))
            if ach_bullets: elements.append(ListFlowable(ach_bullets, bulletType='bullet', start='•', leftIndent=25))
        elements.append(Spacer(1, 6))

    elements.append(Paragraph("Education", style_section))
    edu_bullets = []
    for edu in data.get('education', []):
        text = f"{edu.get('degree', '')}, {edu.get('college', '')}"
        edu_bullets.append(ListItem(Paragraph(clean(text), style_normal), leftIndent=0))
    if edu_bullets: elements.append(ListFlowable(edu_bullets, bulletType='bullet', start='•', leftIndent=15))

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()

# --- 9. UI LAYER ---
st.set_page_config(page_title=PAGE_TITLE, layout="wide", page_icon="🚀", initial_sidebar_state="expanded")

st.markdown("""
<style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .block-container {padding-top: 1.5rem;}
    div.stButton > button:first-child {
        border-radius: 6px;
        font-weight: 600;
        text-transform: none;
    }
    div[data-testid="stMetricValue"] {
        font-size: 1.8rem;
    }
</style>
""", unsafe_allow_html=True)

if 'data' not in st.session_state: st.session_state['data'] = None
if 'saved_base' not in st.session_state: st.session_state['saved_base'] = ""
if 'saved_jd' not in st.session_state: st.session_state['saved_jd'] = ""
if 'cover_letter' not in st.session_state: st.session_state['cover_letter'] = None

with st.sidebar:
    st.header("⚙️ Configuration")
    google_key = st.text_input("Google API Key", type="password")
    groq_key = st.text_input("Groq API Key", type="password")
    
    st.divider()
    if st.button("🗑️ Reset Application", use_container_width=True):
        st.session_state['data'] = None
        st.session_state['saved_base'] = ""
        st.session_state['saved_jd'] = ""
        st.session_state['cover_letter'] = None
        st.rerun()
    st.caption(f"Astra Engine v2.5 (Pydantic)")

if not st.session_state['data']:
    st.markdown(f"<h1 style='text-align: center;'>{PAGE_TITLE}</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #666;'>Enterprise-Grade Resume Architecture & Optimization</p>", unsafe_allow_html=True)
    st.divider()
    
    c1, c2 = st.columns(2)
    with c1:
        st.subheader("📋 Base Resume")
        base = st.text_area("Paste your current resume here", st.session_state['saved_base'], height=400, label_visibility="collapsed")
    with c2:
        st.subheader("💼 Job Description")
        jd = st.text_area("Paste the JD here", st.session_state['saved_jd'], height=400, label_visibility="collapsed")
    
    if st.button("✨ Architect My Application", type="primary", use_container_width=True):
        if google_key and groq_key and base and jd:
            st.session_state['saved_base'] = base
            st.session_state['saved_jd'] = jd
            
            with st.spinner("Analyzing keywords, optimizing structure, and generating narrative..."):
                data = analyze_and_generate(google_key, groq_key, base, jd)
                if "error" in data: st.error(data['error'])
                else: 
                    st.session_state['data'] = data
                    st.rerun()
        else:
            st.warning("Please provide API Keys and both Resume/JD text.")

else:
    # --- DASHBOARD VIEW ---
    data = st.session_state['data']
    
    # Top Metrics Bar
    c1, c2, c3 = st.columns([1, 4, 1])
    with c2:
        st.markdown(f"## 🎯 Target: {data.get('target_company', 'Company')}")
    with c3:
        st.metric("ATS Match", f"{data.get('ats_score', 0)}%")
        # Display reason even if score is 0 so user knows WHY
        if data.get('ats_score') == 0:
            st.caption(f"Status: {data.get('ats_reason', 'N/A')}")

    # Professional Tabs
    tab_edit, tab_export, tab_cover = st.tabs(["📝 Content Editor", "🚀 Export Documents", "✍️ Cover Letter Strategy"])

    with tab_edit:
        with st.form("edit_form"):
            st.subheader("Candidate Details")
            c1, c2, c3 = st.columns(3)
            data['candidate_name'] = c1.text_input("Full Name", to_text_block(data.get('candidate_name')))
            data['candidate_title'] = c2.text_input("Target Title", to_text_block(data.get('candidate_title')))
            data['contact_info'] = c3.text_input("Contact String", to_text_block(data.get('contact_info')))
            
            data['summary'] = st.text_area("Professional Summary", to_text_block(data.get('summary')), height=120)
            
            st.subheader("Skills & Technologies")
            skills = data.get('skills', {})
            new_skills = {}
            s_cols = st.columns(2)
            for i, (k, v) in enumerate(skills.items()):
                col = s_cols[i % 2]
                new_val = col.text_area(k, to_text_block(v), key=f"skill_{i}", height=80)
                new_skills[k] = new_val.replace('\n', ', ')
            data['skills'] = new_skills
            
            st.subheader("Professional Experience")
            for i, role in enumerate(data.get('experience', [])):
                with st.expander(f"{role.get('role_title', 'Role')} @ {role.get('company', 'Company')}"):
                    c1, c2 = st.columns(2)
                    role['role_title'] = c1.text_input("Role Title", to_text_block(role.get('role_title')), key=f"job_title_{i}")
                    role['company'] = c2.text_input("Company", to_text_block(role.get('company')), key=f"job_comp_{i}")
                    c3, c4 = st.columns(2)
                    role['dates'] = c3.text_input("Dates", to_text_block(role.get('dates')), key=f"job_dates_{i}")
                    role['location'] = c4.text_input("Location", to_text_block(role.get('location')), key=f"job_loc_{i}")
                    
                    role['responsibilities'] = st.text_area("Responsibilities (Bullet Points)", to_text_block(role.get('responsibilities')), height=200, key=f"resp_{i}")
                    role['achievements'] = st.text_area("Key Achievements", to_text_block(role.get('achievements')), height=100, key=f"ach_{i}")

            st.subheader("Education")
            for i, edu in enumerate(data.get('education', [])):
                c1, c2 = st.columns(2)
                edu['degree'] = c1.text_input("Degree", to_text_block(edu.get('degree')), key=f"edu_deg_{i}")
                edu['college'] = c2.text_input("Institution", to_text_block(edu.get('college')), key=f"edu_col_{i}")

            if st.form_submit_button("💾 Save Revisions", type="primary"):
                st.session_state['data'] = data
                st.success("Resume updated successfully!")
                st.rerun()

    with tab_export:
        st.subheader("📥 Download Package")
        c_name = data.get('candidate_name', 'Candidate')
        default_company = data.get('target_company', 'Company')
        target_company = st.text_input("Target Company Name (for file labeling)", default_company)
        
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', c_name.strip().replace(' ', '_'))
        safe_company = re.sub(r'[^a-zA-Z0-9_-]', '_', target_company.strip())
        final_filename = f"{safe_name}_{safe_company}"
        
        c1, c2 = st.columns(2)
        
        doc = create_doc(data)
        bio = io.BytesIO()
        doc.save(bio)
        c1.download_button(
            label="📄 Download Word Doc (Editable)",
            data=bio.getvalue(),
            file_name=f"{final_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
        
        try:
            pdf_data = create_pdf(data)
            c2.download_button(
                label="📕 Download PDF (Submission Ready)",
                data=pdf_data,
                file_name=f"{final_filename}.pdf",
                mime="application/pdf",
                type="secondary",
                use_container_width=True
            )
        except Exception as e: c2.error(f"PDF Error: {e}")

    with tab_cover:
        st.subheader("✍️ Strategic Cover Letter")
        st.info("This tool drafts a narrative-driven cover letter focusing on the specific pain points found in the Job Description.")
        
        if st.button("✨ Draft Cover Letter", type="primary"):
            if google_key and st.session_state['saved_jd']:
                with st.spinner("Analyzing JD pain points and drafting narrative..."):
                    cl_text = generate_cover_letter(google_key, data, st.session_state['saved_jd'])
                    st.session_state['cover_letter'] = cl_text
            else:
                st.warning("Please ensure API Key and Job Description are present.")

        if st.session_state['cover_letter']:
            st.text_area("Preview (Editable)", st.session_state['cover_letter'], height=400)
            
            cl_doc = create_cover_letter_doc(st.session_state['cover_letter'], data)
            bio_cl = io.BytesIO()
            cl_doc.save(bio_cl)
            
            st.download_button(
                label="📄 Download Cover Letter (.docx)",
                data=bio_cl.getvalue(),
                file_name=f"Cover_Letter_{final_filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary"
            )

    st.divider()
    c3, c4 = st.columns(2)
    if c3.button("♻️ Re-Optimize (Redo)", use_container_width=True):
        if st.session_state['saved_base'] and st.session_state['saved_jd']:
            with st.spinner("Re-Architecting..."):
                data = analyze_and_generate(google_key, groq_key, st.session_state['saved_base'], st.session_state['saved_jd'])
                if "error" in data: st.error(data['error'])
                else: 
                    st.session_state['data'] = data
                    st.rerun()
                    
    if c4.button("New Application (Keep Resume)", use_container_width=True):
        st.session_state['data'] = None
        st.session_state['saved_jd'] = "" 
        st.session_state['cover_letter'] = None
        st.rerun()
